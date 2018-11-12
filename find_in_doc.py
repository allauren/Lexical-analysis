from docx import Document
from nltk.stem import PorterStemmer
from nltk.tokenize import sent_tokenize, word_tokenize
import re
from docx.enum.text import WD_COLOR_INDEX
from docx import Document
from docx.text.run import Font, Run
from docx.shared import RGBColor


def findWholeWord(w):
    return re.compile(r'\b({0})\b'.format(w), flags=re.IGNORECASE).search

def tokenizer(line):
    ps = PorterStemmer()
    if type(line) == str :
        word_list = word_tokenize(line)
    else :
        word_list = line
    token = []
    for w in word_list:
        token.append([ps.stem(w), w])
    return token

def check_words(file, words):
    document = Document(file)
    newdoc = Document()
    print(words)
    twords = tokenizer(words)
    for paragraph in document.paragraphs :
        tsentence = tokenizer(paragraph.text)
        nparagraph = newdoc.add_paragraph()
        font = nparagraph.add_run().font
        nparagraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
        nparagraph.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
        nparagraph.paragraph_format.right_indent = paragraph.paragraph_format.right_indent
        nparagraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
        nparagraph.style.name = paragraph.style.name
        for word in tsentence :
            found = 0
            for fword in twords :
                if fword[0]== word[0]:
                    found = 1
            if found :
                font.color.rgb = RGBColor(0x3f, 0x2c, 0x36)
                nparagraph.add_run(word[1] + ' ').bold = True
                font.highlight_color = WD_COLOR_INDEX.TURQUOISE
            else :
                nparagraph.add_run(word[1] + ' ').bold = False
    newdoc.save('newdoc.docx')